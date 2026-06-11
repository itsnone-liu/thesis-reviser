# -*- coding: utf-8 -*-
"""
设计类论文修改程序 - 深度重构版
读取原设计类论文 → 分析诊断 → 重写生成 → AI图片插入DOCX
"""

import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import re
import json
import time
import shutil
import uuid
import argparse
import sys
import base64
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from web import (
    call_llm,
    _init_s, _wt, add_c, add_t, chart_to_bytes,
    extract_charts, extract_tables, extract_drawings_from_text,
    get_drawing_rules_by_type, generate_outline,
    SJ_WORD_LIMITS, SJ_CHAPTERS,
    tasks_db, is_system_busy, system_lock
)

# ================== [ 图像生成API配置 ] ==================
IMAGE_API_KEY = "sk-d4jR7aFAJUNPVN5hcdX8l1bXWa59MLK7gvDCgTkcYaVnvgki"
IMAGE_API_URL = "https://api.allmhub.com/v1/images/generations"
IMAGE_MODEL = "gpt-image-2"


# ================== [ 文档解析 ] ==================
def extract_docx_text(path: str) -> str:
    """使用 python-docx 读取 .docx 文件的所有段落文本，保留换行，过滤空段落"""
    try:
        doc = Document(path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"读取文档失败 {path}: {e}")
        return ""


# ================== [ 论文画像分析 ] ==================
def analyze_original_design_paper(text: str) -> dict:
    """调用 LLM 分析设计类论文全文，提取研究画像"""
    prompt = f"""你是一名设计类学术论文分析专家。请对以下论文全文进行深度分析，输出严格JSON格式（不要markdown代码块，不要任何其他文字）。

论文全文：
{text[:8000]}

请输出以下字段的JSON：
{{
  "title": "论文标题",
  "major": "设计专业（如服装设计、产品设计、室内设计、数字媒体设计等）",
  "design_type": "设计专业类型",
  "design_object": "设计对象",
  "context": "使用场景/背景",
  "target_users": ["用户群体1", "用户群体2"],
  "core_problems": ["问题1", "问题2", "问题3"],
  "design_strategies": ["策略1", "策略2", "策略3"],
  "design_elements": ["要素1", "要素2", "要素3"],
  "outline": {{
    "绪论": "本章核心内容摘要（100字内）",
    "理论基础": "...",
    "问题发现与分析": "...",
    "设计策略与方案": "...",
    "总结与反思": "..."
  }},
  "theories": ["理论1", "理论2"],
  "drawings_mentioned": [
    {{"type": "效果图", "title": "...", "description": "..."}}
  ],
  "data_quality_note": "数据/图纸质量评价"
}}

只输出JSON，严禁输出其他内容。"""

    for attempt in range(3):
        response = call_llm(prompt, max_tokens=3000)
        if not response:
            time.sleep(1)
            continue
        try:
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                result = json.loads(json_match.group())
                # 补齐缺失字段
                defaults = {
                    "title": text[:50].replace('\n', ' ').strip() if text else "未识别标题",
                    "major": "设计学",
                    "design_type": "设计学",
                    "design_object": "某设计对象",
                    "context": "使用场景",
                    "target_users": ["目标用户"],
                    "core_problems": ["问题待识别", "问题待识别", "问题待识别"],
                    "design_strategies": ["策略待识别", "策略待识别", "策略待识别"],
                    "design_elements": ["要素1", "要素2", "要素3"],
                    "outline": {name: "" for name, _ in SJ_CHAPTERS},
                    "theories": ["理论待识别"],
                    "drawings_mentioned": [],
                    "data_quality_note": "LLM解析失败，使用默认画像"
                }
                for key, val in defaults.items():
                    if key not in result:
                        result[key] = val
                if "outline" in result:
                    for name, _ in SJ_CHAPTERS:
                        if name not in result["outline"]:
                            result["outline"][name] = ""
                else:
                    result["outline"] = {name: "" for name, _ in SJ_CHAPTERS}
                return result
        except Exception as e:
            print(f"解析论文分析结果失败（尝试{attempt+1}/3）：{e}")
            time.sleep(1)
            continue

    # 兜底默认
    return {
        "title": text[:50].replace('\n', ' ').strip() if text else "未识别标题",
        "major": "设计学",
        "design_type": "设计学",
        "design_object": "某设计对象",
        "context": "使用场景",
        "target_users": ["目标用户"],
        "core_problems": ["问题待识别", "问题待识别", "问题待识别"],
        "design_strategies": ["策略待识别", "策略待识别", "策略待识别"],
        "design_elements": ["要素1", "要素2", "要素3"],
        "outline": {name: "" for name, _ in SJ_CHAPTERS},
        "theories": ["理论待识别"],
        "drawings_mentioned": [],
        "data_quality_note": "LLM解析失败，使用默认画像"
    }


# ================== [ 深度诊断与重构 ] ==================
def diagnose_and_reconstruct_design(original_profile: dict) -> dict:
    """调用 LLM 基于设计类论文画像进行深度诊断，输出修订计划"""
    prompt = f"""你是一名资深设计类学术论文评审与修改专家。请基于以下论文画像进行深度诊断，并输出修订方案。

论文画像：
{json.dumps(original_profile, ensure_ascii=False, indent=2)}

请输出严格JSON格式（不要markdown代码块，不要任何其他文字）：
{{
  "diagnosis_report": {{
    "overall_score": "1-10分",
    "overall_comment": "总体评价（200字）",
    "issues": [
      {{
        "chapter": "涉及章节",
        "severity": "高/中/低",
        "issue_type": "理论错误/逻辑混乱/结构失衡/图纸不足/表述问题/其他",
        "description": "问题描述",
        "fix_direction": "修正方向"
      }}
    ]
  }},
  "revised_profile": {{
    "title": "优化后的标题",
    "design_object": "设计对象",
    "design_type": "设计专业",
    "context": "使用场景",
    "target_users": ["用户1", "用户2"],
    "core_problems": ["修正后问题1", "问题2", "问题3"],
    "design_strategies": ["策略1", "策略2", "策略3"],
    "design_elements": ["要素1", "要素2", "要素3"]
  }},
  "revised_outline": {{
    "chapter1_intro": {{"1.1": "...", "1.2": "...", "1.3": "..."}},
    "chapter2_theory": {{"2.1": "...", "2.2": "...", "2.3": "..."}},
    "chapter3_problem": {{"3.1": "...", "3.2": "...", "3.3": "..."}},
    "chapter4_design": {{"4.1": "...", "4.2": "...", "4.3": "...", "4.4": "...", "4.5": "..."}},
    "chapter5_conclusion": {{"5.1": "...", "5.2": "...", "5.3": "..."}}
  }},
  "chapter_revisions": {{
    "绪论": {{"issues": [], "fix_requirements": [], "key_points": []}},
    "理论基础": {{"issues": [], "fix_requirements": [], "key_points": []}},
    "问题发现与分析": {{"issues": [], "fix_requirements": [], "key_points": []}},
    "设计策略与方案": {{"issues": [], "fix_requirements": ["必须含至少3个<drawing/>标签，包含效果图和设计图"], "key_points": []}},
    "总结与反思": {{"issues": [], "fix_requirements": [], "key_points": []}}
  }},
  "special_notes": "生成新论文时的特殊注意事项"
}}

注意：
- chapter_revisions 中的每个章节都要有 issues、fix_requirements、key_points 三个字段
- 如果某章节无问题，请留空列表[]
- 设计策略与方案章的 fix_requirements 必须强调图纸要求
- 只输出JSON，严禁输出其他内容。"""

    for attempt in range(3):
        response = call_llm(prompt, max_tokens=4000)
        if not response:
            time.sleep(1)
            continue
        try:
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                result = json.loads(json_match.group())
                # 补齐缺失字段
                if "diagnosis_report" not in result:
                    result["diagnosis_report"] = {"overall_score": "5", "overall_comment": "默认评价", "issues": []}
                if "revised_profile" not in result:
                    result["revised_profile"] = {
                        "title": original_profile.get("title", ""),
                        "design_object": original_profile.get("design_object", ""),
                        "design_type": original_profile.get("design_type", ""),
                        "context": original_profile.get("context", ""),
                        "target_users": original_profile.get("target_users", []),
                        "core_problems": original_profile.get("core_problems", []),
                        "design_strategies": original_profile.get("design_strategies", []),
                        "design_elements": original_profile.get("design_elements", [])
                    }
                if "revised_outline" not in result:
                    result["revised_outline"] = {}
                if "chapter_revisions" not in result:
                    result["chapter_revisions"] = {}
                for name, _ in SJ_CHAPTERS:
                    if name not in result["chapter_revisions"]:
                        result["chapter_revisions"][name] = {"issues": [], "fix_requirements": [], "key_points": []}
                    else:
                        for key in ["issues", "fix_requirements", "key_points"]:
                            if key not in result["chapter_revisions"][name]:
                                result["chapter_revisions"][name][key] = []
                if "special_notes" not in result:
                    result["special_notes"] = ""
                return result
        except Exception as e:
            print(f"解析诊断结果失败（尝试{attempt+1}/3）：{e}")
            time.sleep(1)
            continue

    # 保守兜底
    return {
        "diagnosis_report": {
            "overall_score": "5",
            "overall_comment": "由于LLM解析失败，采用保守默认诊断。建议人工复核论文内容。",
            "issues": []
        },
        "revised_profile": {
            "title": original_profile.get("title", ""),
            "design_object": original_profile.get("design_object", ""),
            "design_type": original_profile.get("design_type", ""),
            "context": original_profile.get("context", ""),
            "target_users": original_profile.get("target_users", []),
            "core_problems": original_profile.get("core_problems", []),
            "design_strategies": original_profile.get("design_strategies", []),
            "design_elements": original_profile.get("design_elements", [])
        },
        "revised_outline": {},
        "chapter_revisions": {
            name: {"issues": [], "fix_requirements": [], "key_points": []}
            for name, _ in SJ_CHAPTERS
        },
        "special_notes": "LLM诊断解析失败，请人工检查原始论文并制定修改策略。"
    }


# ================== [ 诊断报告持久化 ] ==================
def save_diagnosis_report(revision_plan: dict, output_path: str):
    """将 revision_plan['diagnosis_report'] 格式化为易读文本报告并写入文件"""
    diag = revision_plan.get("diagnosis_report", {})
    lines = []
    lines.append("=" * 50)
    lines.append("【论文诊断报告】")
    lines.append("=" * 50)
    lines.append("")
    lines.append(f"总体评分：{diag.get('overall_score', 'N/A')}")
    lines.append("")
    lines.append("【总体评价】")
    lines.append(diag.get("overall_comment", "暂无"))
    lines.append("")
    issues = diag.get("issues", [])
    if issues:
        lines.append("【问题清单】")
        for i, issue in enumerate(issues, 1):
            lines.append(f"  问题{i}：")
            lines.append(f"    涉及章节：{issue.get('chapter', 'N/A')}")
            lines.append(f"    严重程度：{issue.get('severity', 'N/A')}")
            lines.append(f"    问题类型：{issue.get('issue_type', 'N/A')}")
            lines.append(f"    问题描述：{issue.get('description', 'N/A')}")
            lines.append(f"    修正方向：{issue.get('fix_direction', 'N/A')}")
            lines.append("")
    else:
        lines.append("【问题清单】")
        lines.append("  未发现明显问题（或LLM未返回问题列表）")
        lines.append("")
    lines.append("=" * 50)

    dir_name = os.path.dirname(output_path)
    if dir_name:
        os.makedirs(dir_name, exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(lines))


# ================== [ 画像融合 ] ==================
def merge_into_standard_design_profile(revised_profile: dict) -> dict:
    """将 revised_profile 转换为兼容原 web.py 设计类的标准 profile"""
    return {
        "title": revised_profile.get("title", ""),
        "design_object": revised_profile.get("design_object", ""),
        "design_type": revised_profile.get("design_type", ""),
        "context": revised_profile.get("context", ""),
        "target_users": revised_profile.get("target_users", []),
        "core_problems": revised_profile.get("core_problems", []),
        "design_strategies": revised_profile.get("design_strategies", []),
        "design_elements": revised_profile.get("design_elements", []),
        "object_type": revised_profile.get("object_type", "设计对象"),
    }


# ================== [ 重写章节生成 ] ==================
def revise_gen_design_chapter(name, num, profile, outline, chapter_rev, update):
    """统一的设计类各章重写生成函数"""
    design_type = profile.get('design_type', '')
    design_object = profile.get('design_object', '')

    if name == "绪论":
        prompt = f"""写设计类论文【第一章 绪论】，研究对象：{design_object}
设计专业：{design_type}

【严格约束】：
1. 不要写"第一章"，直接写内容
2. 小标题格式：1.1、1.2、1.3
3. 字数：800字左右

【内容要求】：
1.1 研究背景：简单交代为什么做这个设计，当前设计领域的问题或机遇
1.2 研究问题：明确要解决什么问题（1-2句话）
1.3 研究意义：为什么重要，对设计实践或理论的价值

只输出正文，不要markdown。
"""
    elif name == "理论基础":
        prompt = f"""写设计类论文【第二章 相关理论基础】，研究对象：{design_object}
设计专业：{design_type}

【严格约束】：
1. 不要写"第二章"，直接写内容
2. 小标题格式：2.1、2.2、2.3
3. 字数：1000字左右

【内容要求】：
2.1 核心理论一介绍（与{design_type}相关的核心理论）
2.2 核心理论二介绍（与设计类型匹配的支撑理论）
2.3 理论与本研究的关联：说明如何用这些理论指导{design_object}的设计

只输出正文，不要markdown。
"""
    elif name == "问题发现与分析":
        prompt = f"""写设计类论文【第三章 问题发现与分析】，研究对象：{design_object}
设计专业：{design_type}

【严格约束】：
1. 不要写"第三章"，直接写内容
2. 小标题格式：3.1、3.2、3.3
3. 字数：1500字左右

【内容要求】：
3.1 {design_object}概况：介绍基本情况、背景、现状
3.2 现状调研与分析：通过调研发现当前存在的问题
3.3 核心问题归纳：归纳{len(profile.get('core_problems', [3]))}个核心问题

【重要】：
- 只写现状和问题，不写解决方案
- 问题要有具体表现

只输出正文，不要markdown。
"""
    elif name == "设计策略与方案":
        strategies = profile.get('design_strategies', ['策略一', '策略二', '策略三'])
        design_chapter = outline.get('chapter4_design', {})
        subsection_titles = []
        for key in sorted(design_chapter.keys()):
            if key.startswith('4.') and key != '4.1':
                subsection_titles.append(f"{key} {design_chapter[key]}")

        subsections_str = ""
        for i, title in enumerate(subsection_titles):
            subsections_str += f"{title}\n"
            if i == 0:
                subsections_str += f'   <drawing id="{i+1}" type="设计图" title="{design_object}设计图" description="展示设计方案"/>\n'
            elif i == 1:
                subsections_str += f'   <drawing id="{i+1}" type="效果图" title="{design_object}效果图" description="展示设计效果"/>\n'
            else:
                subsections_str += f'   <drawing id="{i+1}" type="细节图" title="设计细节图" description="展示设计细节"/>\n'

        drawing_rules = get_drawing_rules_by_type(design_type)

        prompt = f"""写设计类论文【第四章 设计策略与方案】，研究对象：{design_object}
设计专业：{design_type}

【严格约束】：
1. 不要写"第四章"，直接写内容
2. 小标题严格按照以下格式
3. 字数：2500字左右

【小标题结构】：
4.1 设计策略（必须包含以下策略并与第三章问题对应）
   - 策略一：{strategies[0] if strategies else '策略一'} —— 解决什么问题
   - 策略二：{strategies[1] if len(strategies) > 1 else '策略二'} —— 解决什么问题
   - 策略三：{strategies[2] if len(strategies) > 2 else '策略三'} —— 解决什么问题

{subsections_str}

【图纸要求】：
{drawing_rules}

【写作要求】：
- 每个策略要详细说明如何对应解决第三章的问题
- 每个设计小标题下写300-500字的设计说明
- 设计说明要专业、具体，符合{design_type}专业术语

只输出正文，不要markdown。
"""
    else:  # 总结与反思
        prompt = f"""写设计类论文【第五章 总结与反思】，研究对象：{design_object}

【严格约束】：
1. 不要写"第五章"，直接写内容
2. 小标题格式：5.1、5.2、5.3
3. 字数：800字左右

【内容要求】：
5.1 主要成果：完成了什么设计，解决了什么问题
5.2 不足之处：设计中的局限，如时间有限、图纸不够细致等
5.3 改进方向：未来可以如何完善

只输出正文，不要markdown。
"""

    # 追加修改要求
    issues = chapter_rev.get("issues", []) if chapter_rev else []
    fix_reqs = chapter_rev.get("fix_requirements", []) if chapter_rev else []
    key_points = chapter_rev.get("key_points", []) if chapter_rev else []

    if issues or fix_reqs or key_points:
        prompt += "\n【修改要求】（必须严格遵守）：\n"
        if issues:
            prompt += "需要修复的问题：\n" + "\n".join([f"- {i}" for i in issues]) + "\n"
        if fix_reqs:
            prompt += "具体修改要求：\n" + "\n".join([f"- {f}" for f in fix_reqs]) + "\n"
        if key_points:
            prompt += "必须体现的关键点：\n" + "\n".join([f"- {k}" for k in key_points]) + "\n"

    update(f"正在生成第{num}章：{name}...", int(10 + num * 14))
    return call_llm(prompt, max_tokens=4000)


# ================== [ 图片提示词生成 ] ==================
def generate_image_prompt(drawing_info: dict, context_text: str, design_type: str) -> str:
    """调用LLM生成高质量英文图像生成提示词"""
    prompt = f"""你是一名专业的AI图像提示词工程师。
请根据以下设计图纸信息，生成一段高质量的英文图像生成提示词（100-300词），用于指导AI生成设计类图纸。

【设计专业】：{design_type}
【图纸类型】：{drawing_info['type']}
【图纸标题】：{drawing_info['title']}
【图纸描述】：{drawing_info['description']}
【上下文内容】：{context_text[:500]}

要求：
1. 必须是英文
2. 描述要详细、专业，包含风格、构图、色彩、材质、光影等要素
3. 适合{design_type}专业的设计图纸
4. 直接输出提示词，不要任何解释
"""
    return call_llm(prompt, max_tokens=500)


# ================== [ 图片生成 ] ==================
def generate_image(prompt: str, output_path: str) -> bool:
    """调用中转站图像生成API并保存图片"""
    try:
        headers = {
            "Authorization": f"Bearer {IMAGE_API_KEY}",
            "Content-Type": "application/json"
        }
        data = {
            "model": IMAGE_MODEL,
            "prompt": prompt,
            "n": 1,
            "size": "1024x1024",
            "quality": "high",
            "output_format": "png"
        }
        response = requests.post(IMAGE_API_URL, headers=headers, json=data, timeout=180)
        if response.status_code == 200:
            b64_data = response.json()["data"][0]["b64_json"]
            img_bytes = base64.b64decode(b64_data)
            with open(output_path, "wb") as f:
                f.write(img_bytes)
            return True
        else:
            print(f"图片生成失败：{response.status_code} {response.text}")
            return False
    except Exception as e:
        print(f"图片生成异常：{e}")
        return False


# ================== [ 图纸位置提取 ] ==================
def extract_drawings_with_positions(text: str) -> list:
    """从文本中提取图纸标记，并返回在文本中的位置"""
    drawings = []
    pattern = re.compile(
        r'<drawing\s+id=["\']?(\d+)["\']?\s+type=["\']([^"\']*)["\']\s+title=["\']([^"\']*)["\']\s+description=["\']([^"\']*)["\']\s*/?>'
    )
    for match in pattern.finditer(text):
        drawings.append({
            "id": int(match.group(1)),
            "type": match.group(2),
            "title": match.group(3),
            "description": match.group(4),
            "start_pos": match.start(),
            "end_pos": match.end()
        })
    return drawings


# ================== [ DOCX图纸插入 ] ==================
def add_drawing_image(doc, drawing, image_path, dn):
    """向DOCX中插入设计图纸图片及图注（全局编号：图1、图2...）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.add_run().add_picture(image_path, width=Inches(5.5))
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    r = p2.add_run(f"图{dn} {drawing['title']}")
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
    r.font.size = Pt(10.5)


# ================== [ 带图片的DOCX排版 ] ==================
def design_txt_to_docx_safe(txt_path, docx_path, drawing_images, update):
    """复制 txt_to_docx_safe 核心逻辑，扩展支持 drawing 图片插入"""
    if drawing_images is None:
        drawing_images = {}
    update("正在读取TXT...", 10)
    with open(txt_path, "r", encoding="utf-8") as f:
        full_text = f.read()
    update("正在提取图表与图纸数据...", 30)
    charts = extract_charts(full_text)
    tables = extract_tables(full_text)
    drawings = extract_drawings_with_positions(full_text)
    els = (
        [(c["start_pos"], c["end_pos"], "chart", c) for c in charts] +
        [(t["start_pos"], t["end_pos"], "table", t) for t in tables] +
        [(d["start_pos"], d["end_pos"], "drawing", d) for d in drawings]
    )
    els.sort(key=lambda x: x[0])
    parts, le = [], 0
    for s, e, t, d in els:
        if s > le:
            parts.append(("text", full_text[le:s]))
        parts.append((t, d))
        le = e
    if le < len(full_text):
        parts.append(("text", full_text[le:]))
    update("正在生成Word排版...", 60)
    doc = Document()
    _init_s(doc)
    cn, tn, dn = 0, 0, 0
    for pt, ct in parts:
        if pt == "chart":
            cn += 1
            try:
                add_c(doc, ct, chart_to_bytes(ct), cn)
            except Exception as e:
                print(f"生成图表失败: {e}")
        elif pt == "table":
            tn += 1
            try:
                add_t(doc, ct, tn)
            except Exception as e:
                print(f"生成表格失败: {e}")
        elif pt == "drawing":
            dn += 1
            try:
                img_path = drawing_images.get(ct["id"])
                if img_path and os.path.exists(img_path):
                    add_drawing_image(doc, ct, img_path, dn)
                else:
                    print(f"图纸图片缺失 id={ct['id']}, 跳过")
            except Exception as e:
                print(f"插入图纸失败: {e}")
        else:
            _wt(doc, ct.strip())
    doc.save(docx_path)
    update("DOCX生成完毕！", 100)


# ================== [ 设计类重写管道 ] ==================
def run_revise_design_pipeline(task_id, profile, revision_plan):
    """设计类论文修改生成管道"""
    folder = f"output/{task_id}"
    os.makedirs(folder, exist_ok=True)

    def update(msg, prog):
        tasks_db[task_id].update({"msg": msg, "progress": prog})
        print(f"[{prog}%] {msg}")

    try:
        update("正在融合修改画像与大纲...", 5)
        outline = revision_plan.get("revised_outline", {})
        if not outline:
            outline = generate_outline(profile)

        update("正在生成摘要...", 10)
        special_notes = revision_plan.get("special_notes", "")
        abstract_prompt = f"""写设计类论文摘要，400字左右。
题目：{profile['title']}
设计对象：{profile['design_object']}
设计专业：{profile['design_type']}
核心问题：{profile.get('core_problems', [])}
设计策略：{profile.get('design_strategies', [])}

【特别注意事项】：
{special_notes}

要求：学术流畅，包含背景、问题、策略、成果、意义。
只输出摘要正文。
"""
        abstract = call_llm(abstract_prompt, max_tokens=800)

        keywords_prompt = f"""根据以下信息生成3-5个设计类论文关键词，用分号隔开：
设计对象：{profile['design_object']}
设计专业：{profile['design_type']}

只输出关键词，如：服装设计；可持续时尚；系列设计
"""
        keywords = call_llm(keywords_prompt, max_tokens=200).strip()

        chapter_revisions = revision_plan.get("chapter_revisions", {})
        chapters_content = []
        chapter_names = ["绪论", "理论基础", "问题发现与分析", "设计策略与方案", "总结与反思"]

        for name, num in zip(chapter_names, range(1, 6)):
            chapter_rev = chapter_revisions.get(name, {})
            content = revise_gen_design_chapter(name, num, profile, outline, chapter_rev, update)
            chapters_content.append(content)

        update("正在生成参考文献...", 85)
        refs_prompt = f"""生成12-15条设计类论文规范参考文献，类型包括书籍[M]、期刊论文[J]、学位论文[D]。
主题涉及：{profile['design_type']}、{profile['design_object']}相关设计理论。
每条占一行，格式如：[1] 作者. 书名[M]. 出版社, 年份.
"""
        refs_res = call_llm(refs_prompt, max_tokens=1500)
        refs = [r.strip() for r in refs_res.split('\n') if r.strip() and (r.strip()[0].isdigit() or r.strip().startswith('['))]
        refs = refs[:15]

        # 合并论文
        txt = f"{profile['title']}\n\n摘要\n{abstract}\n\n关键词\n{keywords}\n\n目录\n\n"
        for i, (content, name) in enumerate(zip(chapters_content, chapter_names)):
            txt += f"第{i+1}章 {name}\n{content.strip()}\n\n"

        txt += "参考文献\n" + "\n".join(refs)

        txt_path = os.path.join(folder, "00_完整论文.txt")
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(txt)

        # 提取图纸（用于后续图片生成）
        design_chapter_text = chapters_content[3] if len(chapters_content) > 3 else ""
        drawings = extract_drawings_from_text(design_chapter_text)

        update("TXT生成完毕！", 100)
        tasks_db[task_id].update({
            "status": "completed",
            "files": [txt_path],
            "drawings": drawings
        })
    except Exception as e:
        tasks_db[task_id].update({"status": "error", "msg": f"生成失败: {str(e)}"})
        print(f"生成失败: {e}")


# ================== [ 主入口 ] ==================
def main():
    parser = argparse.ArgumentParser(description="设计类论文修改程序 - 深度重构")
    parser.add_argument("--input", "-i", required=True, help="输入的DOCX文件路径")
    parser.add_argument("--output", "-o", default="./output", help="输出目录（默认./output）")
    args = parser.parse_args()

    if not os.path.isfile(args.input):
        print(f"错误：输入文件不存在 {args.input}")
        return

    os.makedirs(args.output, exist_ok=True)

    print("=" * 50)
    print("【设计类论文修改程序】启动")
    print("=" * 50)

    print("\n步骤1/4：提取DOCX文本...")
    original_text = extract_docx_text(args.input)
    if not original_text:
        print("错误：未能从DOCX中提取文本，程序终止。")
        return
    print(f"  提取成功，共 {len(original_text)} 字符")

    print("\n步骤2/4：分析原文画像...")
    analysis_result = analyze_original_design_paper(original_text)
    print(f"  识别标题：{analysis_result.get('title', 'N/A')}")
    print(f"  识别专业：{analysis_result.get('design_type', 'N/A')}")
    print(f"  识别对象：{analysis_result.get('design_object', 'N/A')}")

    print("\n步骤3/4：诊断与重构...")
    revision_plan = diagnose_and_reconstruct_design(analysis_result)
    report_path = os.path.join(args.output, "诊断报告.txt")
    save_diagnosis_report(revision_plan, report_path)
    print(f"  诊断报告已保存：{report_path}")
    diag = revision_plan.get("diagnosis_report", {})
    print(f"  总体评分：{diag.get('overall_score', 'N/A')}")
    print(f"  共发现 {len(diag.get('issues', []))} 个问题")

    print("\n步骤4/4：生成修改版论文...")
    revised_profile = revision_plan.get("revised_profile", {})
    profile = merge_into_standard_design_profile(revised_profile)

    task_id = str(uuid.uuid4())
    tasks_db[task_id] = {"status": "running", "progress": 0, "msg": "启动", "files": []}

    run_revise_design_pipeline(task_id, profile, revision_plan)

    if tasks_db[task_id].get("status") != "completed":
        print(f"错误：论文生成失败，{tasks_db[task_id].get('msg', '未知错误')}")
        return

    folder = f"output/{task_id}"
    txt_path = os.path.join(folder, "00_完整论文.txt")
    drawings = tasks_db[task_id].get("drawings", [])

    print(f"\n  发现 {len(drawings)} 张图纸需要生成...")
    drawing_images = {}
    for d in drawings:
        print(f"  生成图片: {d['title']} ({d['type']})...")
        with open(txt_path, "r", encoding="utf-8") as f:
            context_text = f.read()
        img_prompt = generate_image_prompt(d, context_text, profile.get("design_type", ""))
        img_path = os.path.join(folder, f"drawing_{d['id']}.png")
        if generate_image(img_prompt, img_path):
            drawing_images[d["id"]] = img_path
            print(f"    成功: {img_path}")
        else:
            print(f"    失败，跳过")

    print("\n  生成DOCX排版...")
    docx_path = os.path.join(folder, "论文_修改版.docx")

    def update_fn(msg, prog):
        print(f"  [DOCX {prog}%] {msg}")

    design_txt_to_docx_safe(txt_path, docx_path, drawing_images, update_fn)
    finalize_docx(docx_path, update_fn)

    final_docx = os.path.join(args.output, "论文_修改版.docx")
    shutil.copy(docx_path, final_docx)
    print(f"\n{'=' * 50}")
    print(f"全部完成！")
    print(f"诊断报告：{report_path}")
    print(f"修改版论文：{final_docx}")
    print(f"{'=' * 50}")


if __name__ == "__main__":
    main()


def finalize_docx(docx_path, update):
    """分节符修正 + PDF 真实页码 + 目录更新 + 页脚"""
    update("正在处理分节符和页码...", 92)
    doc = Document(docx_path)
    paragraphs = doc.paragraphs
    body = doc.element.body

    toc_idx = -1
    for i, p in enumerate(paragraphs):
        if p.text.strip() == '目录':
            toc_idx = i
            break

    chapter1_idx = -1
    for i, p in enumerate(paragraphs):
        t = p.text.strip()
        if re.match(r'^第1章\s', t) and i > toc_idx + 3 and '\t' not in t:
            chapter1_idx = i
            break

    if chapter1_idx < 0:
        print("[跳过] 未找到第1章")
        return

    toc_end_idx = chapter1_idx - 1
    while toc_end_idx > toc_idx:
        if paragraphs[toc_end_idx].text.strip():
            break
        toc_end_idx -= 1

    toc_headings = []
    toc_range_start = toc_idx + 1 if toc_idx >= 0 else 0
    for i in range(toc_range_start, chapter1_idx):
        t = paragraphs[i].text.strip()
        if t:
            heading = t.split('\t')[0].strip() if '\t' in t else t
            toc_headings.append(heading)

    # 保存临时副本 → 转 PDF → 算页码
    tag = uuid.uuid4().hex[:8]
    temp_docx = docx_path.replace('.docx', f'_temp_{tag}.docx')
    doc.save(temp_docx)

    update("正在用 LibreOffice 渲染 PDF 计算页码...", 95)
    pdf_path = os.path.join(os.path.dirname(temp_docx) or '.', f'_temp_{tag}.pdf')
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf',
                    temp_docx, '--outdir', os.path.dirname(pdf_path)],
                   capture_output=True, text=True, timeout=120)
    expected = temp_docx.replace('.docx', '.pdf')
    if os.path.exists(expected):
        os.rename(expected, pdf_path)

    heading_abs_pages = {}
    if os.path.exists(pdf_path):
        result = subprocess.run(['pdfinfo', pdf_path], capture_output=True, text=True)
        num_pages = 0
        for line in result.stdout.split('\n'):
            if line.startswith('Pages'):
                num_pages = int(line.split(':')[1].strip())
                break

        # 找目录页并跳过
        toc_pdf_page = -1
        for pg in range(1, num_pages + 1):
            r = subprocess.run(['pdftotext', '-f', str(pg), '-l', str(pg), pdf_path, '-'],
                               capture_output=True, text=True)
            lines = [l.strip() for l in r.stdout.split('\n') if l.strip()]
            if any(l == '目录' for l in lines):
                toc_pdf_page = pg
                break

        start_pg = toc_pdf_page + 1 if toc_pdf_page >= 0 else 1
        for pg in range(start_pg, num_pages + 1):
            r = subprocess.run(['pdftotext', '-f', str(pg), '-l', str(pg), pdf_path, '-'],
                               capture_output=True, text=True)
            text_flat = re.sub(r'\s+', '', r.stdout)
            for h in toc_headings:
                if h in heading_abs_pages:
                    continue
                h_flat = re.sub(r'\s+', '', h)
                if h_flat in text_flat:
                    heading_abs_pages[h] = pg

        ch1_heading = next((h for h in toc_headings if h.startswith('第1章')), None)
        ch1_page = heading_abs_pages.get(ch1_heading, 1)
        offset = ch1_page - 1

        # 更新目录页码
        update("正在更新目录页码...", 97)
        for i in range(toc_range_start, chapter1_idx):
            p = paragraphs[i]
            t = p.text.strip()
            if not t:
                continue
            heading_text = t.split('\t')[0].strip() if '\t' in t else t
            if heading_text in heading_abs_pages:
                sec_page = heading_abs_pages[heading_text] - offset
                if sec_page >= 1:
                    p.clear()
                    tab_stops = p.paragraph_format.tab_stops
                    tab_stops.add_tab_stop(Cm(16), alignment=WD_TAB_ALIGNMENT.RIGHT,
                                           leader=WD_TAB_LEADER.DOTS)
                    r = p.add_run(heading_text)
                    set_run_font(r, "宋体", 14)
                    p.alignment = 0
                    p.add_run("\t")
                    rp = p.add_run(str(sec_page))
                    set_run_font(rp, "宋体", 14)

        if os.path.exists(pdf_path):
            os.remove(pdf_path)

    # 分节符
    last_sect_pr = None
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'sectPr':
            last_sect_pr = child

    if last_sect_pr is not None:
        template = deepcopy(last_sect_pr)
        body.remove(last_sect_pr)

        toc_end_elem = paragraphs[toc_end_idx]._element
        toc_end_pPr = toc_end_elem.find(qn('w:pPr'))
        if toc_end_pPr is None:
            toc_end_pPr = OxmlElement('w:pPr')
            toc_end_elem.insert(0, toc_end_pPr)
        sPr0 = OxmlElement('w:sectPr')
        for c in template:
            tag = c.tag.split('}')[-1] if '}' in c.tag else c.tag
            if tag in ('pgSz', 'pgMar'):
                sPr0.append(deepcopy(c))
        toc_end_pPr.append(sPr0)

        ts = OxmlElement('w:sectPr')
        for c in template:
            tag = c.tag.split('}')[-1] if '}' in c.tag else c.tag
            if tag in ('pgSz', 'pgMar'):
                ts.append(deepcopy(c))
        pgnt = OxmlElement('w:pgNumType')
        pgnt.set(qn('w:start'), '1')
        ts.append(pgnt)
        body.append(ts)

    # 移除第1章的 pageBreakBefore
    ch1_elem = paragraphs[chapter1_idx]._element
    ch1_pPr = ch1_elem.find(qn('w:pPr'))
    if ch1_pPr is not None:
        for pb in ch1_pPr.findall(qn('w:pageBreakBefore')):
            ch1_pPr.remove(pb)

    # 页脚
    try:
        sec0 = doc.sections[0]
        sec0.different_first_page_header_footer = True
        for fn in ['footer', 'even_page_footer', 'first_page_footer']:
            try:
                f = getattr(sec0, fn)
                f.is_linked_to_previous = False
                for pf in f.paragraphs:
                    pf.clear()
            except Exception:
                pass
    except Exception:
        pass

    try:
        if len(doc.sections) > 1:
            sec1 = doc.sections[1]
            footer = sec1.footer
            footer.is_linked_to_previous = False
            for pf in footer.paragraphs:
                pf.clear()
            pf = footer.paragraphs[0]
            pf.alignment = 1
            run = pf.add_run()
            fc1 = OxmlElement('w:fldChar')
            fc1.set(qn('w:fldCharType'), 'begin')
            run._element.append(fc1)
            it = OxmlElement('w:instrText')
            it.set(qn('xml:space'), 'preserve')
            it.text = ' PAGE '
            run._element.append(it)
            fc2 = OxmlElement('w:fldChar')
            fc2.set(qn('w:fldCharType'), 'end')
            run._element.append(fc2)
    except Exception as e:
        print(f"页脚处理异常: {e}")

    doc.save(docx_path)
    if os.path.exists(temp_docx):
        os.remove(temp_docx)
    update("最终排版完成", 100)



    finalize_docx(docx_path, update_fn)