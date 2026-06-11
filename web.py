# -*- coding: utf-8 -*-
"""
论文生成 Web 服务【经管类 + 设计类 整合版】
启动：pip install fastapi uvicorn python-docx matplotlib numpy requests python-multipart
      python web.py
"""

import os, re, io, json, time, uuid, threading, shutil
import numpy as np
from typing import Dict
import requests
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
plt.rcParams["font.sans-serif"] = ["SimHei"]
plt.rcParams["axes.unicode_minus"] = False
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from fastapi import FastAPI, UploadFile, File, BackgroundTasks, Request
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from pydantic import BaseModel

# ================== [ 配置区 ] ==================
DEEPSEEK_API_KEY = "sk-5c8f324baf804186b32a9869461c01c2" # 必须改成你的！
DEEPSEEK_BASE_URL = "https://api.deepseek.com/v1"
COLORS = ["#5B9BD5", "#ED7D31", "#A5A5A5", "#FFC000", "#4472C4", "#70AD47", "#264478"]

# 经管类配置
MG_WORD_LIMITS = {"引言": 1000, "国内外研究现状": 1500, "现状分析": 2500, "问题与原因分析": 3000, "解决方案": 3000, "结论": 500}
MG_SOFT_MAX = {"引言": 1200, "国内外研究现状": 1800, "现状分析": 3500, "问题与原因分析": 4500, "解决方案": 4000, "结论": 800}
MG_CHAPTERS = [("引言", 1), ("国内外研究现状", 2), ("现状分析", 3), ("问题与原因分析", 4), ("解决方案", 5), ("结论", 6)]

# 设计类配置
SJ_WORD_LIMITS = {"绪论": 800, "理论基础": 1000, "问题发现与分析": 1500, "设计策略与方案": 2500, "总结与反思": 800}
SJ_CHAPTERS = [("绪论", 1), ("理论基础", 2), ("问题发现与分析", 3), ("设计策略与方案", 4), ("总结与反思", 5)]

# ================== [ 单线程排队锁 ] ==================
tasks_db = {}
is_system_busy = False
system_lock = threading.Lock()

# ================== [ LLM 核心调用 ] ==================
def call_llm(prompt: str, max_tokens: int = 4096, retry: int = 3) -> str:
    headers = {"Authorization": f"Bearer {DEEPSEEK_API_KEY}", "Content-Type": "application/json"}
    data = {"model": "deepseek-chat", "messages": [{"role": "user", "content": prompt}], "temperature": 0.7, "max_tokens": max_tokens}
    for attempt in range(retry):
        try:
            response = requests.post(f"{DEEPSEEK_BASE_URL}/chat/completions", headers=headers, json=data, timeout=180)
            if response.status_code == 200: return response.json()["choices"][0]["message"]["content"].strip()
            else:
                print(f"API调用失败（尝试{attempt+1}/{retry}）：{response.status_code}")
                if attempt < retry - 1: time.sleep(2)
        except Exception as e:
            print(f"API调用异常（尝试{attempt+1}/{retry}）：{str(e)}")
            if attempt < retry - 1: time.sleep(2)
    print("API调用失败，返回空内容")
    return ""

def build_chart_rule():
    return """
【图表与表格输出规范 - 必须严格遵守】

一、图表（chart）规范：
<chart id="1" title="图表标题" type="bar或line或pie" x="标签1,标签2,标签3" y="数值1,数值2,数值3" unit="单位" data_source="来源"/>

关键要求：
1. y属性必须是数字，用逗号分隔，如 y="100,150,200"
2. x属性是标签，用逗号分隔，如 x="2020,2021,2022"
3. 严禁在y属性中放文字说明！

二、表格（table）规范：
<table id="1" title="表格标题" header="列1,列2,列3" rows="行名1,行名2,行名3" data="行1数据;行2数据;行3数据" data_source="来源"/>

关键要求：
1. rows属性必须是行名称列表，用逗号分隔，如 rows="营业收入,净利润,总资产"
2. 严禁在rows属性中只写数字（如rows="3"是错误的）！
3. data属性中，多行数据必须用分号;分隔
4. data属性中，同一行的多列数据必须用竖线|分隔
5. 正确示例：data="100|50|200;150|75|250;200|100|300"
   - 表示3行数据，每行3列
   - 第一行：100, 50, 200
   - 第二行：150, 75, 250
   - 第三行：200, 100, 300

三、常见错误（严禁）：
❌ 错误1：rows="3" （这是行数，不是行名）
✅ 正确：rows="营业收入,净利润,总资产"

❌ 错误2：data="586400|492000|15.8%,213500|240000|-11.0%" （用逗号分隔行）
✅ 正确：data="586400|492000|15.8%;213500|240000|-11.0%"

❌ 错误3：y="预算成本,实际成本" （y属性必须是数字）
✅ 正确：y="100,150,200"

四、完整示例：
<table id="1" title="设备费用对比" header="设备名称,实际费用,标准费用,差异率" rows="数控中心,液压机,车床" data="586400|492000|15.8%;213500|240000|-11.0%;98500|126000|-21.8%" data_source="公司财务部"/>"""

# ================== [ 经管类核心逻辑 ] ==================
def generate_profile_web(major, title, company):
    prompt = f"""你是经管类论文研究设计专家。 
【专业】：{major} 
【论文题目】：{title} 
【研究对象】：{company} 
请分析对象，生成严格JSON格式的研究画像： 
{{ 
  "company": "{company}", 
  "org_type": "类型", 
  "industry": "行业", 
  "core_problems": ["问题1","问题2","问题3"], 
  "data_hints": ["数据1","数据2","数据3","数据4"] 
}} 
只输出JSON。"""
    response = call_llm(prompt, max_tokens=1200)
    try:
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            profile = json.loads(json_match.group())
            profile["major"] = major
            profile["title"] = title
            return profile
    except:
        pass
    return {
        "company": company, "org_type": "企业", "industry": "制造业", 
        "core_problems": ["问题1","问题2","问题3"], 
        "data_hints": ["数据1","数据2","数据3"], 
        "major": major, "title": title
    }

def gen_chapter(name, num, profile, update):
    update(f"正在生成第{num}章 {name}...", int((num-1)/6 * 80))
    if name == "引言": sp = "\n- 禁止使用任何标题、编号、markdown。用2-3段写背景和意义。"
    elif name == "国内外研究现状": sp = f"\n- 禁止使用任何标题、编号、markdown。纯段落叙述，不谈{profile['company']}。"
    elif name == "现状分析": sp = "\n- 小标题：3.1、3.2、3.3。必须含至少2个<chart/>和1个<table/>。"
    elif name == "问题与原因分析": sp = "\n- 小标题：4.1、4.2、4.3。必须含至少1个<chart/>和1个<table/>。"
    elif name == "解决方案": sp = "\n- 小标题：5.1、5.2、5.3。必须含至少1个<chart/>"
    elif name == "结论": sp = "\n- 300-500字总结，无编号。"
    else: sp = ""
    cr = build_chart_rule() if name in ["现状分析", "问题与原因分析", "解决方案"] else ""
    hd = "禁止任何标题格式" if name in ["引言", "国内外研究现状", "结论"] else f"小标题：{num}.1 / {num}.2格式"
    prompt = f"""写《{profile['title']}》第{num}章 {name}。字数：{MG_WORD_LIMITS[name]}内。
格式铁律：1.{hd} 2.无大标题 3.无markdown符号 4.无加粗 5.纯文本正文{sp}
{cr} 对象：{profile['company']} 行业：{profile['industry']} 只输出正文。"""
    text = call_llm(prompt)
    return re.sub(r'^第[一二三四五六\d]+章.*?\n', '', text).strip()

def run_txt_pipeline(task_id: str, profile: Dict):
    global is_system_busy
    folder = f"output/{task_id}"
    os.makedirs(folder, exist_ok=True)
    def update(msg, prog): tasks_db[task_id].update({"msg": msg, "progress": prog})
    try:
        update("正在生成摘要...", 5)
        abstract = call_llm(f"写300字摘要：{profile['title']}，对象{profile['company']}。", 600)
        txt = f"{profile['title']}\n\n摘要\n{abstract}\n\n关键词\n{profile['industry']}；{profile['company']}\n\n目录\n\n"
        for name, num in MG_CHAPTERS:
            txt += f"第{num}章 {name}\n{gen_chapter(name, num, profile, update)}\n\n"
        update("正在生成参考文献...", 85)
        refs = call_llm(f"围绕{profile['title']}，生成12条规范参考文献[M][J][D]。", 1500)
        txt += "参考文献\n" + "\n".join([r.strip() for r in refs.split('\n') if r.strip().startswith('[')][:12])
        
        txt_path = os.path.join(folder, "00_完整论文.txt")
        with open(txt_path, 'w', encoding='utf-8') as f: f.write(txt)
        update("TXT生成完毕！", 100)
        tasks_db[task_id].update({"status": "completed", "files": [txt_path]})
    except Exception as e:
        tasks_db[task_id].update({"status": "error", "msg": f"生成失败: {str(e)}"})
    finally:
        with system_lock: is_system_busy = False

# ================== [ 设计类核心逻辑 ] ==================
def generate_design_profile(design_type: str, title: str, target_name: str, force_variation: bool = False) -> Dict:
    """生成设计类研究画像 - 支持所有设计专业"""
    variation_hint = ""
    if force_variation:
        variation_hint = "\n必须生成完全不同的设计问题和策略。"
    
    prompt = f"""
你是设计类论文研究设计专家。
【设计专业】：{design_type}
【论文题目】：{title}
【设计对象】：{target_name}
{variation_hint}

请分析设计对象，生成严格JSON格式的研究画像：
{{
    "design_object": "{target_name}",
    "design_type": "{design_type}",
    "object_type": "对象类型（根据专业判断）",
    "context": "设计背景/使用场景",
    "target_users": ["用户群体1", "用户群体2"],
    "core_problems": ["问题1", "问题2", "问题3"],
    "design_strategies": ["策略1", "策略2", "策略3"],
    "design_elements": ["设计要素1", "设计要素2", "设计要素3"]
}}

只输出JSON，不要其他内容。
"""
    response = call_llm(prompt, max_tokens=1200)
    try:
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            profile = json.loads(json_match.group())
            profile["design_type"] = design_type
            profile["title"] = title
            return profile
    except:
        pass
    
    return {
        "design_object": target_name,
        "design_type": design_type,
        "object_type": "设计对象",
        "context": "使用场景",
        "target_users": ["目标用户"],
        "core_problems": ["问题1", "问题2", "问题3"],
        "design_strategies": ["策略一", "策略二", "策略三"],
        "design_elements": ["要素1", "要素2", "要素3"],
        "title": title
    }

def generate_outline(profile: Dict) -> Dict:
    """动态生成大纲 - LLM根据专业/题目/对象自主生成合适的章节结构"""
    design_type = profile.get('design_type', '')
    design_object = profile.get('design_object', '')
    title = profile.get('title', '')
    core_problems = profile.get('core_problems', [])
    design_strategies = profile.get('design_strategies', [])
    
    prompt = f"""
你是设计类论文结构专家。请根据以下信息，为这篇设计类毕业论文生成一个合适的章节大纲。

【设计专业】：{design_type}
【论文题目】：{title}
【设计对象】：{design_object}
【核心问题】：{core_problems}
【初步设计策略】：{design_strategies}

【任务要求】：
1. 根据专业特点，生成合理的5章结构（绪论、理论基础、问题发现、设计策略与方案、总结反思）
2. 第4章"设计策略与方案"的子标题要符合该专业的设计流程和术语
3. 不同专业应有不同侧重点：
   - 服装与服饰设计：可包含设计定位、款式设计、面料选择、色彩搭配、结构工艺、系列设计等
   - 数字媒体设计：可包含需求分析、交互设计、视觉设计、动效设计、技术实现、用户体验测试等
   - 产品设计：可包含造型设计、功能结构、CMF设计、人机交互等
   - 室内设计：可包含空间布局、功能分区、材质软装、灯光设计等
   - 景观设计：可包含总体布局、分区设计、植物配置、小品设施等
   - 建筑设计：可包含总体布局、单体设计、空间序列、构造设计等
   - 视觉传达：可包含标志设计、应用系统、色彩字体、设计规范等

【输出格式】：
必须输出严格的JSON，格式如下：
{{
    "chapter1_intro": {{
        "1.1": "研究背景",
        "1.2": "研究问题",
        "1.3": "研究意义"
    }},
    "chapter2_theory": {{
        "2.1": "核心理论一介绍",
        "2.2": "核心理论二介绍",
        "2.3": "理论与本设计的关联"
    }},
    "chapter3_problem": {{
        "3.1": "项目/对象概况",
        "3.2": "现状调研与分析",
        "3.3": "核心问题归纳"
    }},
    "chapter4_design": {{
        "4.1": "设计策略",
        "4.2": "【根据专业特点生成】",
        "4.3": "【根据专业特点生成】",
        "4.4": "【根据专业特点生成】",
        "4.5": "【根据专业特点生成】"
    }},
    "chapter5_conclusion": {{
        "5.1": "主要成果",
        "5.2": "不足之处",
        "5.3": "改进方向"
    }}
}}

【重要】：
- 第4章必须有3-5个子标题（4.2到4.5之间）
- 子标题要符合{design_type}专业的设计流程和术语
- 只输出JSON，不要其他任何内容
"""
    response = call_llm(prompt, max_tokens=2000)
    try:
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            outline = json.loads(json_match.group())
            if 'chapter4_design' in outline:
                design_chapter = outline['chapter4_design']
                subsections = [k for k in design_chapter.keys() if k.startswith('4.') and k != '4.1']
                if len(subsections) < 3:
                    for i in range(len(subsections) + 1, 5):
                        design_chapter[f"4.{i}"] = f"设计细化{i-1}"
            return outline
    except Exception as e:
        print(f"大纲解析失败：{e}")
    
    return get_fallback_outline(design_type, design_object)

def get_fallback_outline(design_type: str, design_object: str) -> Dict:
    """备用大纲"""
    if '服装' in design_type or '服饰' in design_type:
        return {
            "chapter1_intro": {"1.1": "研究背景", "1.2": "研究问题", "1.3": "研究意义"},
            "chapter2_theory": {"2.1": "服装设计理论", "2.2": "色彩与面料理论", "2.3": "理论与本设计的关联"},
            "chapter3_problem": {"3.1": "设计现状调研", "3.2": "市场与用户分析", "3.3": "核心问题归纳"},
            "chapter4_design": {
                "4.1": "设计策略", "4.2": "设计定位与灵感来源",
                "4.3": "款式与结构设计", "4.4": "面料与色彩设计", "4.5": "系列成品展示"
            },
            "chapter5_conclusion": {"5.1": "主要成果", "5.2": "不足之处", "5.3": "改进方向"}
        }
    elif '数字媒体' in design_type or '数媒' in design_type or '交互' in design_type:
        return {
            "chapter1_intro": {"1.1": "研究背景", "1.2": "研究问题", "1.3": "研究意义"},
            "chapter2_theory": {"2.1": "交互设计理论", "2.2": "用户体验理论", "2.3": "理论与本设计的关联"},
            "chapter3_problem": {"3.1": "项目需求分析", "3.2": "竞品与用户调研", "3.3": "核心问题归纳"},
            "chapter4_design": {
                "4.1": "设计策略", "4.2": "信息架构与交互设计",
                "4.3": "视觉界面设计", "4.4": "动效与音效设计", "4.5": "技术实现与测试"
            },
            "chapter5_conclusion": {"5.1": "主要成果", "5.2": "不足之处", "5.3": "改进方向"}
        }
    elif '产品' in design_type or '工业' in design_type:
        return {
            "chapter1_intro": {"1.1": "研究背景", "1.2": "研究问题", "1.3": "研究意义"},
            "chapter2_theory": {"2.1": "人机工程学", "2.2": "设计心理学", "2.3": "理论与本设计的关联"},
            "chapter3_problem": {"3.1": "产品现状调研", "3.2": "用户需求分析", "3.3": "核心问题归纳"},
            "chapter4_design": {
                "4.1": "设计策略", "4.2": "产品造型设计",
                "4.3": "功能与结构设计", "4.4": "CMF设计", "4.5": "用户体验设计"
            },
            "chapter5_conclusion": {"5.1": "主要成果", "5.2": "不足之处", "5.3": "改进方向"}
        }
    else:
        return {
            "chapter1_intro": {"1.1": "研究背景", "1.2": "研究问题", "1.3": "研究意义"},
            "chapter2_theory": {"2.1": "核心理论介绍", "2.2": "相关支撑理论", "2.3": "理论与本设计的关联"},
            "chapter3_problem": {"3.1": f"{design_object}概况", "3.2": "现状调研分析", "3.3": "核心问题归纳"},
            "chapter4_design": {
                "4.1": "设计策略", "4.2": "总体设计方案",
                "4.3": "方案细化一", "4.4": "方案细化二", "4.5": "专项设计"
            },
            "chapter5_conclusion": {"5.1": "主要成果", "5.2": "不足之处", "5.3": "改进方向"}
        }

def get_drawing_rules_by_type(design_type: str) -> str:
    """根据专业类型返回图纸标记要求"""
    if '服装' in design_type or '服饰' in design_type:
        return """
【图纸标记格式】：
<drawing id="1" type="款式图" title="服装款式设计图" description="展示服装廓形和结构线"/>
<drawing id="2" type="效果图" title="服装效果图" description="展示穿着效果和整体造型"/>
<drawing id="3" type="面料小样" title="面料选择图" description="展示面料材质和色彩"/>
<drawing id="4" type="结构图" title="服装结构图" description="展示版型和工艺细节"/>

要求：必须包含至少1张款式图 + 1张效果图
"""
    elif '数字媒体' in design_type or '数媒' in design_type or '交互' in design_type:
        return """
【图纸标记格式】：
<drawing id="1" type="流程图" title="交互流程图" description="展示用户操作流程"/>
<drawing id="2" type="线框图" title="界面线框图" description="展示信息架构和布局"/>
<drawing id="3" type="高保真图" title="界面效果图" description="展示视觉设计效果"/>
<drawing id="4" type="动效图" title="动效设计图" description="展示交互动效"/>

要求：必须包含至少1张流程图/线框图 + 1张效果图
"""
    elif '产品' in design_type or '工业' in design_type:
        return """
【图纸标记格式】：
<drawing id="1" type="效果图" title="产品效果图" description="展示产品外观"/>
<drawing id="2" type="三视图" title="产品三视图" description="展示尺寸和比例"/>
<drawing id="3" type="爆炸图" title="结构爆炸图" description="展示内部结构"/>
<drawing id="4" type="场景图" title="使用场景图" description="展示使用状态"/>

要求：必须包含至少1张效果图 + 1张三视图
"""
    elif '室内' in design_type or '空间' in design_type:
        return """
【图纸标记格式】：
<drawing id="1" type="平面图" title="平面布置图" description="展示功能分区"/>
<drawing id="2" type="效果图" title="空间效果图" description="展示设计效果"/>
<drawing id="3" type="立面图" title="主要立面图" description="展示墙面设计"/>

要求：必须包含至少1张平面图 + 1张效果图
"""
    elif '景观' in design_type or '环境' in design_type:
        return """
【图纸标记格式】：
<drawing id="1" type="总平面图" title="景观总平面图" description="展示整体布局"/>
<drawing id="2" type="效果图" title="景观效果图" description="展示设计效果"/>
<drawing id="3" type="剖立面图" title="剖立面图" description="展示竖向设计"/>

要求：必须包含至少1张总平面图 + 1张效果图
"""
    elif '建筑' in design_type:
        return """
【图纸标记格式】：
<drawing id="1" type="总平面图" title="建筑总平面图" description="展示建筑与环境"/>
<drawing id="2" type="效果图" title="建筑效果图" description="展示建筑外观"/>
<drawing id="3" type="平面图" title="建筑平面图" description="展示功能布局"/>

要求：必须包含至少1张总平面图 + 1张效果图
"""
    elif '视觉' in design_type or '平面' in design_type:
        return """
【图纸标记格式】：
<drawing id="1" type="标志设计" title="标志设计图" description="展示标志图形"/>
<drawing id="2" type="应用效果图" title="应用效果图" description="展示应用场景"/>
<drawing id="3" type="规范图" title="设计规范图" description="展示色彩字体规范"/>

要求：必须包含至少1个标志设计 + 1张应用效果图
"""
    else:
        return """
【图纸标记格式】：
<drawing id="1" type="设计图" title="设计方案图" description="展示设计方案"/>
<drawing id="2" type="效果图" title="设计效果图" description="展示设计效果"/>

要求：必须包含至少1张设计图 + 1张效果图
"""

def generate_intro(profile: Dict) -> str:
    """生成绪论章节"""
    prompt = f"""
写设计类论文【第一章 绪论】，研究对象：{profile['design_object']}
设计专业：{profile['design_type']}

【严格约束】：
1. 不要写"第一章"，直接写内容
2. 小标题格式：1.1、1.2、1.3
3. 字数：800字左右

【内容要求】：
1.1 研究背景：简单交代为什么做这个设计，当前设计领域的问题或机遇
1.2 研究问题：明确要解决什么问题（1-2句话）
1.3 研究意义：为什么重要，对设计实践或理论的价值

只输出正文，不要markdown。
"""
    return call_llm(prompt, max_tokens=2000)

def generate_theory(profile: Dict, outline: Dict) -> str:
    """生成理论基础章节"""
    prompt = f"""
写设计类论文【第二章 相关理论基础】，研究对象：{profile['design_object']}
设计专业：{profile['design_type']}

【严格约束】：
1. 不要写"第二章"，直接写内容
2. 小标题格式：2.1、2.2、2.3
3. 字数：1000字左右

【内容要求】：
2.1 核心理论一介绍（与{profile['design_type']}相关的核心理论）
2.2 核心理论二介绍（与设计类型匹配的支撑理论）
2.3 理论与本研究的关联：说明如何用这些理论指导{profile['design_object']}的设计

只输出正文，不要markdown。
"""
    return call_llm(prompt, max_tokens=2500)

def generate_problem_analysis(profile: Dict, outline: Dict) -> str:
    """生成问题发现与分析章节"""
    prompt = f"""
写设计类论文【第三章 问题发现与分析】，研究对象：{profile['design_object']}
设计专业：{profile['design_type']}

【严格约束】：
1. 不要写"第三章"，直接写内容
2. 小标题格式：3.1、3.2、3.3
3. 字数：1500字左右

【内容要求】：
3.1 {profile['design_object']}概况：介绍基本情况、背景、现状
3.2 现状调研与分析：通过调研发现当前存在的问题
3.3 核心问题归纳：归纳{len(profile.get('core_problems', [3]))}个核心问题

【重要】：
- 只写现状和问题，不写解决方案
- 问题要有具体表现

只输出正文，不要markdown。
"""
    return call_llm(prompt, max_tokens=3500)

def generate_design_strategy_and_solution(profile: Dict, outline: Dict) -> str:
    """生成设计策略与方案章节 - 根据大纲动态生成"""
    design_type = profile.get('design_type', '')
    design_object = profile.get('design_object', '')
    strategies = profile.get('design_strategies', ['策略一', '策略二', '策略三'])
    
    design_chapter = outline.get('chapter4_design', {})
    subsection_titles = []
    for key in sorted(design_chapter.keys()):
        if key.startswith('4.') and key != '4.1':
            subsection_titles.append(f"{key} {design_chapter[key]}")
    
    subsections_str = ""
    for i, title in enumerate(subsection_titles):
        subsections_str += f"{title}\n"
        if i == 0:
            subsections_str += f'   <drawing id="{i+1}" type="设计图" title="{design_object}设计图" description="展示设计方案"/>\n'
        elif i == 1:
            subsections_str += f'   <drawing id="{i+1}" type="效果图" title="{design_object}效果图" description="展示设计效果"/>\n'
        else:
            subsections_str += f'   <drawing id="{i+1}" type="细节图" title="设计细节图" description="展示设计细节"/>\n'
    
    drawing_rules = get_drawing_rules_by_type(design_type)
    
    prompt = f"""
写设计类论文【第四章 设计策略与方案】，研究对象：{design_object}
设计专业：{design_type}

【严格约束】：
1. 不要写"第四章"，直接写内容
2. 小标题严格按照以下格式
3. 字数：2500字左右

【小标题结构】：
4.1 设计策略（必须包含以下策略并与第三章问题对应）
   - 策略一：{strategies[0] if strategies else '策略一'} —— 解决什么问题
   - 策略二：{strategies[1] if len(strategies) > 1 else '策略二'} —— 解决什么问题
   - 策略三：{strategies[2] if len(strategies) > 2 else '策略三'} —— 解决什么问题

{subsections_str}

【图纸要求】：
{drawing_rules}

【写作要求】：
- 每个策略要详细说明如何对应解决第三章的问题
- 每个设计小标题下写300-500字的设计说明
- 设计说明要专业、具体，符合{design_type}专业术语

只输出正文，不要markdown。
"""
    return call_llm(prompt, max_tokens=4000)

def generate_conclusion(profile: Dict, outline: Dict) -> str:
    """生成总结与反思章节"""
    prompt = f"""
写设计类论文【第五章 总结与反思】，研究对象：{profile['design_object']}

【严格约束】：
1. 不要写"第五章"，直接写内容
2. 小标题格式：5.1、5.2、5.3
3. 字数：800字左右

【内容要求】：
5.1 主要成果：完成了什么设计，解决了什么问题
5.2 不足之处：设计中的局限，如时间有限、图纸不够细致等
5.3 改进方向：未来可以如何完善

只输出正文，不要markdown。
"""
    return call_llm(prompt, max_tokens=2000)

def generate_abstract(profile: Dict) -> str:
    """生成摘要"""
    prompt = f"""
写设计类论文摘要，400字左右。
题目：{profile['title']}
设计对象：{profile['design_object']}
设计专业：{profile['design_type']}
核心问题：{profile.get('core_problems', [])}
设计策略：{profile.get('design_strategies', [])}

要求：学术流畅，包含背景、问题、策略、成果、意义。
只输出摘要正文。
"""
    return call_llm(prompt, max_tokens=800)

def generate_keywords(profile: Dict) -> str:
    """生成关键词"""
    prompt = f"""
根据以下信息生成3-5个设计类论文关键词，用分号隔开：
设计对象：{profile['design_object']}
设计专业：{profile['design_type']}

只输出关键词，如：服装设计；可持续时尚；系列设计
"""
    keywords = call_llm(prompt, max_tokens=200)
    return keywords.strip()

def generate_references(profile: Dict) -> list:
    """生成参考文献"""
    prompt = f"""
生成12-15条设计类论文规范参考文献，类型包括书籍[M]、期刊论文[J]、学位论文[D]。
主题涉及：{profile['design_type']}、{profile['design_object']}相关设计理论。
每条占一行，格式如：[1] 作者. 书名[M]. 出版社, 年份.
"""
    res = call_llm(prompt, max_tokens=1500)
    refs = [r.strip() for r in res.split('\n') if r.strip() and (r.strip()[0].isdigit() or r.strip().startswith('['))]
    return refs[:15]

def extract_drawings_from_text(text: str) -> list:
    """从文本中提取图纸标记"""
    drawings = []
    pattern = re.compile(
        r'<drawing\s+id=["\']?(\d+)["\']?\s+type=["\']([^"\']*)["\']\s+title=["\']([^"\']*)["\']\s+description=["\']([^"\']*)["\']\s*/?>'
    )
    matches = pattern.findall(text)
    for m in matches:
        drawings.append({
            "id": int(m[0]),
            "type": m[1],
            "title": m[2],
            "description": m[3]
        })
    return drawings

def run_design_txt_pipeline(task_id: str, profile: Dict):
    """设计类论文生成管道"""
    global is_system_busy
    folder = f"output/{task_id}"
    os.makedirs(folder, exist_ok=True)
    def update(msg, prog): tasks_db[task_id].update({"msg": msg, "progress": prog})
    try:
        update("正在生成设计大纲...", 5)
        outline = generate_outline(profile)
        
        update("正在生成摘要...", 10)
        abstract = generate_abstract(profile)
        keywords = generate_keywords(profile)
        
        update("正在生成第一章：绪论...", 20)
        intro = generate_intro(profile)
        
        update("正在生成第二章：理论基础...", 35)
        theory = generate_theory(profile, outline)
        
        update("正在生成第三章：问题发现与分析...", 50)
        problem = generate_problem_analysis(profile, outline)
        
        update("正在生成第四章：设计策略与方案...", 65)
        design = generate_design_strategy_and_solution(profile, outline)
        
        update("正在生成第五章：总结与反思...", 80)
        conclusion = generate_conclusion(profile, outline)
        
        update("正在生成参考文献...", 90)
        refs = generate_references(profile)
        
        # 合并论文
        txt = f"{profile['title']}\n\n摘要\n{abstract}\n\n关键词\n{keywords}\n\n目录\n\n"
        chapters = [intro, theory, problem, design, conclusion]
        chapter_names = ["绪论", "理论基础", "问题发现与分析", "设计策略与方案", "总结与反思"]
        for i, (content, name) in enumerate(zip(chapters, chapter_names)):
            txt += f"第{i+1}章 {name}\n{content.strip()}\n\n"
        
        txt += "参考文献\n" + "\n".join(refs[:15])
        txt += "\n\n附录\n设计图纸目录\n"
        
        # 提取图纸
        drawings = extract_drawings_from_text(design)
        if drawings:
            txt += "\n".join([f"【图纸】{d['type']}: {d['title']}" for d in drawings])
        
        txt_path = os.path.join(folder, "00_完整论文.txt")
        with open(txt_path, 'w', encoding='utf-8') as f: f.write(txt)
        update("TXT生成完毕！", 100)
        tasks_db[task_id].update({"status": "completed", "files": [txt_path]})
    except Exception as e:
        tasks_db[task_id].update({"status": "error", "msg": f"生成失败: {str(e)}"})
    finally:
        with system_lock: is_system_busy = False

# ================== [ 智能标签解析系统 ] ==================

def _extract_attr(tag_text, attr_name, default=""):
    """从标签文本中提取属性值，支持多种引号和空格"""
    patterns = [
        rf'{attr_name}\s*=\s*"([^"]*)"',
        rf"{attr_name}\s*=\s*'([^']*)",
        rf'{attr_name}\s*=\s*([^\s>]+)',
    ]
    for pattern in patterns:
        match = re.search(pattern, tag_text, re.I)
        if match:
            return match.group(1).strip()
    return default

def _smart_split_values(text, expected_count=None):
    """智能分割数值字符串"""
    if not text:
        return []
    delimiters = [',', ';', '|', ' ', '\t']
    best_result = []
    for delimiter in delimiters:
        if delimiter in text:
            parts = [p.strip() for p in text.split(delimiter) if p.strip()]
            if len(parts) > len(best_result):
                best_result = parts
    if not best_result:
        best_result = [text.strip()]
    return best_result

def _try_parse_number(s):
    """尝试将字符串解析为数字，支持各种格式"""
    if not s:
        return None
    cleaned = s.strip()
    try:
        return float(cleaned)
    except ValueError:
        pass
    patterns = [
        r'[-+]?\d{1,3}(?:,\d{3})+\.?\d*',
        r'[-+]?\d+\.?\d*',
    ]
    for pattern in patterns:
        match = re.search(pattern, cleaned)
        if match:
            num_str = match.group().replace(',', '')
            try:
                return float(num_str)
            except ValueError:
                continue
    return None

def _is_pure_number_list(text):
    """检查字符串是否全是数字（可能包含分隔符）"""
    if not text:
        return False
    parts = _smart_split_values(text)
    if not parts:
        return False
    number_count = sum(1 for p in parts if _try_parse_number(p) is not None)
    return number_count >= len(parts) * 0.5

def extract_charts(text):
    """智能提取图表标签"""
    charts = []
    pattern = re.compile(r'<chart\b([^>]*)(?:/>|>(.*?)</chart>)', re.I | re.S)
    for match in pattern.finditer(text):
        try:
            tag_text = match.group(1)
            inner_text = match.group(2) if match.group(2) else ""
            chart_id = _extract_attr(tag_text, 'id', '0')
            title = _extract_attr(tag_text, 'title', '未命名图表')
            chart_type = _extract_attr(tag_text, 'type', 'bar').lower()
            x_str = _extract_attr(tag_text, 'x', '')
            y_str = _extract_attr(tag_text, 'y', '')
            unit = _extract_attr(tag_text, 'unit', '')
            data_source = _extract_attr(tag_text, 'data[-_]?source', '')
            if not y_str and inner_text:
                y_str = inner_text.strip()
            x_labels = _smart_split_values(x_str)
            y_series = []
            if y_str:
                if ';' in y_str:
                    series_list = [s.strip() for s in y_str.split(';') if s.strip()]
                else:
                    series_list = [y_str.strip()]
                for series in series_list:
                    values = []
                    parts = _smart_split_values(series)
                    for part in parts:
                        num = _try_parse_number(part)
                        if num is not None:
                            values.append(num)
                    if values:
                        y_series.append(values)
            if not y_series or not any(y_series):
                print(f"警告: 图表 '{title}' 没有有效的数值数据，跳过")
                continue
            if all(all(v == 0 for v in series) for series in y_series):
                print(f"警告: 图表 '{title}' 的数据全为0，可能是文字而非数字，跳过")
                continue
            valid_types = ['bar', 'line', 'pie']
            if chart_type not in valid_types:
                chart_type = 'bar'
            if chart_type == 'pie' and len(y_series) > 1:
                y_series = [y_series[0]]
            expected_len = len(y_series[0]) if y_series else 0
            if x_labels and len(x_labels) != expected_len:
                if len(x_labels) < expected_len:
                    while len(x_labels) < expected_len:
                        x_labels.append(f"项{len(x_labels)+1}")
                else:
                    x_labels = x_labels[:expected_len]
            elif not x_labels and expected_len > 0:
                x_labels = [f"项{i+1}" for i in range(expected_len)]
            charts.append({
                "id": int(chart_id) if chart_id.isdigit() else len(charts) + 1,
                "title": title,
                "type": chart_type,
                "x": x_labels,
                "y_series": y_series,
                "y": y_series[0] if y_series else [],
                "unit": unit,
                "data_source": data_source,
                "start_pos": match.start(),
                "end_pos": match.end()
            })
        except Exception as e:
            print(f"解析图表失败: {e}")
            continue
    return charts

def extract_tables(text):
    """智能提取表格标签"""
    tables = []
    pattern = re.compile(r'<table\b([^>]*)(?:/>|>(.*?)</table>)', re.I | re.S)
    for match in pattern.finditer(text):
        try:
            tag_text = match.group(1)
            inner_text = match.group(2) if match.group(2) else ""
            table_id = _extract_attr(tag_text, 'id', '0')
            title = _extract_attr(tag_text, 'title', '未命名表格')
            header_str = _extract_attr(tag_text, 'header', '')
            rows_str = _extract_attr(tag_text, 'rows', '')
            data_str = _extract_attr(tag_text, 'data', '')
            data_source = _extract_attr(tag_text, 'data[-_]?source', '')
            if not data_str and inner_text:
                data_str = inner_text.strip()
            headers = _smart_split_values(header_str)
            header_count = len(headers)
            row_names = []
            if rows_str:
                if rows_str.strip().isdigit():
                    pass
                else:
                    row_names = _smart_split_values(rows_str)
            parsed_rows = []
            if data_str:
                if ';' in data_str:
                    raw_rows = [r.strip() for r in data_str.split(';') if r.strip()]
                elif '|' in data_str and ',' in data_str:
                    comma_parts = [p.strip() for p in data_str.split(',') if p.strip()]
                    if all('|' in p for p in comma_parts):
                        raw_rows = comma_parts
                    else:
                        raw_rows = [data_str.strip()]
                elif '\n' in data_str:
                    raw_rows = [r.strip() for r in data_str.split('\n') if r.strip()]
                else:
                    raw_rows = [data_str.strip()]
                for i, row_data in enumerate(raw_rows):
                    cells = []
                    if '|' in row_data:
                        cells = [c.strip() for c in row_data.split('|')]
                    elif ',' in row_data:
                        cells = [c.strip() for c in row_data.split(',')]
                    else:
                        cells = [row_data.strip()]
                    if i < len(row_names):
                        cells.insert(0, row_names[i])
                    elif row_names:
                        cells.insert(0, f"行{i+1}")
                    parsed_rows.append(cells)
            if not parsed_rows and rows_str and rows_str.strip().isdigit():
                num_rows = int(rows_str.strip())
                for i in range(num_rows):
                    parsed_rows.append([f"行{i+1}"])
            if not parsed_rows:
                print(f"警告: 表格 '{title}' 没有数据行，跳过")
                continue
            max_cols = max(len(row) for row in parsed_rows)
            has_row_names = any(len(row) > len(headers) for row in parsed_rows)
            expected_header_cols = max_cols - (1 if has_row_names else 0)
            while len(headers) < expected_header_cols:
                headers.append(f"列{len(headers)+1}")
            if len(headers) > expected_header_cols:
                headers = headers[:expected_header_cols]
            for row in parsed_rows:
                while len(row) < max_cols:
                    row.append('')
                if len(row) > max_cols:
                    row[:] = row[:max_cols]
            tables.append({
                "id": int(table_id) if table_id.isdigit() else len(tables) + 1,
                "title": title,
                "headers": headers,
                "rows": parsed_rows,
                "data_source": data_source,
                "start_pos": match.start(),
                "end_pos": match.end()
            })
        except Exception as e:
            print(f"解析表格失败: {e}")
            continue
    return tables

# ================== [ CDT 核心逻辑 ] ==================

def chart_to_bytes(c):
    fig,ax=plt.subplots(figsize=(10,5.5))
    ys,x,t=c.get("y_series",[]),c["x"],c["type"]
    if not ys or not x:
        return io.BytesIO()
    if t=="bar":
        if len(ys)==1:
            bs=ax.bar(x,ys[0],color=COLORS[0],width=0.6,edgecolor="w")
            for b,v in zip(bs,ys[0]): ax.text(b.get_x()+b.get_width()/2,b.get_height(),f"{v:g}",ha='center',va='bottom',fontsize=9)
        else:
            n=len(ys);xi=np.arange(len(x));w=0.8/n
            for i,yv in enumerate(ys): ax.bar(xi+(i-n/2+.5)*w,yv,w,label=f"系列{i+1}",color=COLORS[i%len(COLORS)],edgecolor="w")
            ax.set_xticks(xi);ax.set_xticklabels(x,fontsize=9);ax.legend(fontsize=9)
    elif t=="line":
        for i,yv in enumerate(ys): ax.plot(x,yv,marker="o",lw=2,color=COLORS[i%len(COLORS)],ms=7,mfc="w",mew=2,mec=COLORS[i%len(COLORS)]); [ax.text(j,v,f"{v:g}",ha='center',va='bottom',fontsize=8) for j,v in enumerate(yv)]
    elif t=="pie":
        if ys and ys[0]:
            ax.pie(ys[0],labels=x,autopct="%1.1f%%",colors=COLORS[:len(x)],startangle=90,pctdistance=0.75,textprops={"fontsize":10})
    else:
        if ys and ys[0]:
            ax.bar(x,ys[0],color=COLORS[0],width=0.6)
    ax.set_title(c["title"],fontsize=13,fontweight="bold",pad=12)
    if c.get("unit") and t!="pie": ax.set_ylabel(c["unit"],fontsize=11)
    ax.spines["top"].set_visible(False);ax.spines["right"].set_visible(False)
    if t!="pie": ax.tick_params(axis="x",labelsize=9)
    plt.tight_layout()
    buf=io.BytesIO();fig.savefig(buf,format="png",dpi=200,bbox_inches="tight",facecolor="w");plt.close(fig);buf.seek(0)
    return buf

def _is_fake(tb):
    tb=tb.strip()
    if not tb: return True
    if '。' not in tb and any(k in tb for k in ['数据来源','图1','图2','表1','表2','考核维度','财务指标']): return True
    return False

def _init_s(doc):
    n=doc.styles['Normal'];n.font.name="宋体";n._element.rPr.rFonts.set(qn('w:eastAsia'),"宋体");n.font.size=Pt(12)
    for l,cn,sz,b,a in [(1,"黑体",18,True,WD_ALIGN_PARAGRAPH.CENTER),(2,"黑体",14,True,WD_ALIGN_PARAGRAPH.LEFT)]:
        sn=f"Heading {l}"
        if sn in [s.name for s in doc.styles]:
            s=doc.styles[sn];s.font.name=cn;s._element.rPr.rFonts.set(qn('w:eastAsia'),cn);s.font.size=Pt(sz);s.font.bold=b;s.font.color.rgb=RGBColor(0,0,0);s.paragraph_format.alignment=a

def add_c(doc,c,img,cn):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(12)
    r=p.add_run(f"图{cn} {c['title']}");r.font.name="黑体";r._element.rPr.rFonts.set(qn('w:eastAsia'),"黑体");r.font.size=Pt(10.5)
    p2=doc.add_paragraph();p2.alignment=WD_ALIGN_PARAGRAPH.CENTER;p2.add_run().add_picture(img,width=Inches(5.5))
    if c.get("data_source"):
        p3=doc.add_paragraph();p3.alignment=WD_ALIGN_PARAGRAPH.CENTER;p3.paragraph_format.space_after=Pt(12)
        r=p3.add_run(f"数据来源：{c['data_source']}");r.font.size=Pt(9);r.font.color.rgb=RGBColor(128,128,128)

def add_t(doc,t,tn):
    """添加表格到文档，包含数据验证和修复"""
    h,r=t["headers"],t["rows"]
    if not r:
        print(f"警告: 表格 '{t.get('title', '未命名')}' 没有数据行，跳过")
        return
    max_cols = max(len(row) for row in r)
    has_row_names = any(len(row) > len(h) for row in r)
    cc = max_cols
    if len(h) < max_cols - (1 if has_row_names else 0):
        while len(h) < max_cols - (1 if has_row_names else 0):
            h.append(f"列{len(h)+1}")
    if len(h) > max_cols - (1 if has_row_names else 0):
        h = h[:max_cols - (1 if has_row_names else 0)]
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(12)
    rn=p.add_run(f"表{tn} {t['title']}");rn.font.name="黑体";rn._element.rPr.rFonts.set(qn('w:eastAsia'),"黑体");rn.font.size=Pt(10.5)
    tb=doc.add_table(rows=1+len(r),cols=cc,style='Table Grid');tb.alignment=WD_ALIGN_PARAGRAPH.CENTER
    hd=tb.rows[0].cells
    if has_row_names:
        hd[0].text = ""
        for j, hx in enumerate(h):
            if j+1 < cc:
                hd[j+1].text = hx
    else:
        for j, hx in enumerate(h):
            if j < cc:
                hd[j].text = hx
    for ri,rd in enumerate(r):
        row_data = rd[:cc]
        while len(row_data) < cc:
            row_data.append('')
        for ci in range(cc):
            tb.rows[ri+1].cells[ci].text = str(row_data[ci])
    if t.get("data_source"):
        p2=doc.add_paragraph();p2.alignment=WD_ALIGN_PARAGRAPH.CENTER;p2.paragraph_format.space_after=Pt(12)
        r=p2.add_run(f"数据来源：{t['data_source']}");r.font.size=Pt(9);r.font.color.rgb=RGBColor(128,128,128)

def _wt(doc,tb):
    if not tb or _is_fake(tb): return
    for ln in tb.split("\n"):
        ln=ln.strip()
        if not ln: continue
        if re.match(r'^第[一二三四五六\d]+章',ln): doc.add_paragraph(ln,style='Heading 1')
        elif re.match(r'^\d+\.\d+\.\d+',ln):
            p=doc.add_paragraph(ln);r=p.runs[0] if p.runs else p.add_run();r.font.name="黑体";r._element.rPr.rFonts.set(qn('w:eastAsia'),"黑体");r.font.size=Pt(15);r.bold=True;r.font.color.rgb=RGBColor(0,0,0)
        elif re.match(r'^\d+\.\d+',ln): doc.add_paragraph(ln,style='Heading 2')
        elif ln in ['摘要','关键词','目录','参考文献']: doc.add_paragraph(ln,style='Heading 1')
        else: p=doc.add_paragraph(ln);p.paragraph_format.first_line_indent=Cm(0.74);p.paragraph_format.line_spacing=1.5

def txt_to_docx_safe(txt_path, docx_path, update):
    """完全仿照本地双击运行：从硬盘读TXT，写硬盘DOCX"""
    update("正在读取TXT...", 10)
    with open(txt_path, "r", encoding="utf-8") as f: full_text = f.read()
    update("正在提取图表数据...", 30)
    charts, tables = extract_charts(full_text), extract_tables(full_text)
    els = [(c["start_pos"],c["end_pos"],"chart",c) for c in charts] + [(t["start_pos"],t["end_pos"],"table",t) for t in tables]
    els.sort(key=lambda x: x[0])
    parts, le = [], 0
    for s,e,t,d in els:
        if s>le: parts.append(("text",full_text[le:s]))
        parts.append((t,d)); le=e
    if le<len(full_text): parts.append(("text",full_text[le:]))
    update("正在生成Word排版...", 60)
    doc = Document(); _init_s(doc)
    cn, tn = 0, 0
    for pt, ct in parts:
        if pt=="chart":
            cn+=1
            try: add_c(doc, ct, chart_to_bytes(ct), cn)
            except Exception as e: print(f"生成图表失败: {e}")
        elif pt=="table":
            tn+=1
            try: add_t(doc, ct, tn)
            except Exception as e: print(f"生成表格失败: {e}")
        else: _wt(doc, ct.strip())
    doc.save(docx_path)
    update("DOCX生成完毕！", 100)

def run_docx_pipeline(task_id: str, uploaded_txt_path: str):
    global is_system_busy
    folder = f"output/{task_id}"
    os.makedirs(folder, exist_ok=True)
    safe_txt_path = os.path.join(folder, "upload.txt")
    shutil.copy(uploaded_txt_path, safe_txt_path)
    docx_path = os.path.join(folder, "论文_含图表.docx")
    def update(msg, prog): tasks_db[task_id].update({"msg": msg, "progress": prog})
    try:
        txt_to_docx_safe(safe_txt_path, docx_path, update)
        tasks_db[task_id].update({"status": "completed", "files": [docx_path]})
    except Exception as e:
        tasks_db[task_id].update({"status": "error", "msg": f"转换失败: {str(e)}"})
    finally:
        with system_lock: is_system_busy = False

# ================== [ FastAPI 路由 ] ==================
app = FastAPI()

# ---------- 主页面：选择入口 ----------
@app.get("/", response_class=HTMLResponse)
async def index():
    return """
<!DOCTYPE html><html lang="zh-CN"><head><meta charset="UTF-8"><title>论文生成系统</title><script src="https://cdn.tailwindcss.com"></script></head>
<body class="bg-gray-100 p-8">
<div class="max-w-4xl mx-auto space-y-8">
  <h1 class="text-3xl font-bold text-center text-blue-700">🎓 毕业论文生成系统</h1>
  <p class="text-center text-gray-600">请选择论文类型开始生成</p>
  
  <div class="grid grid-cols-2 gap-8">
    <a href="/manage" class="bg-white p-8 rounded-lg shadow border-t-4 border-blue-500 hover:shadow-lg transition-shadow text-center">
      <div class="text-5xl mb-4">📊</div>
      <h2 class="text-xl font-bold text-blue-700 mb-2">经管类论文</h2>
      <p class="text-gray-600">工商管理、财务管理、市场营销等</p>
      <p class="text-sm text-gray-500 mt-2">包含：引言、研究现状、现状分析、问题分析、解决方案、结论</p>
    </a>
    <a href="/design" class="bg-white p-8 rounded-lg shadow border-t-4 border-purple-500 hover:shadow-lg transition-shadow text-center">
      <div class="text-5xl mb-4">🎨</div>
      <h2 class="text-xl font-bold text-purple-700 mb-2">设计类论文</h2>
      <p class="text-gray-600">服装设计、产品设计、室内设计、数字媒体等</p>
      <p class="text-sm text-gray-500 mt-2">包含：绪论、理论基础、问题分析、设计策略与方案、总结反思</p>
    </a>
  </div>
</div>
</body></html>"""

# ---------- 经管类页面 ----------
@app.get("/manage", response_class=HTMLResponse)
async def manage_page():
    return """
<!DOCTYPE html><html lang="zh-CN"><head><meta charset="UTF-8"><title>经管类论文生成</title><script src="https://cdn.tailwindcss.com"></script></head>
<body class="bg-gray-100 p-8">
<div class="max-w-4xl mx-auto space-y-8">
  <div class="flex items-center justify-between">
    <a href="/" class="text-blue-600 hover:underline">← 返回首页</a>
    <h1 class="text-3xl font-bold text-center text-blue-700">📊 经管类毕业论文生成</h1>
    <div></div>
  </div>
  
  <!-- 第一步：生成TXT -->
  <div class="bg-white p-6 rounded-lg shadow border-t-4 border-blue-500">
    <h2 class="text-xl font-bold mb-4">第一步：生成论文 TXT 原文</h2>
    <div class="grid grid-cols-2 gap-4 mb-4">
      <div><label class="block text-sm font-medium mb-1">论文题目</label><input id="t_title" class="w-full border p-2 rounded" value="供应链视角下企业营运资金管理优化研究"></div>
      <div><label class="block text-sm font-medium mb-1">专业</label><input id="t_major" class="w-full border p-2 rounded" value="工商管理"></div>
      <div><label class="block text-sm font-medium mb-1">研究对象</label><input id="t_comp" class="w-full border p-2 rounded" value="K公司"></div>
      <div><label class="block text-sm font-medium mb-1">所属行业</label><input id="t_ind" class="w-full border p-2 rounded" value="制造业"></div>
    </div>
    <button onclick="genProfile()" class="bg-gray-500 text-white px-4 py-2 rounded mb-4">1. 自动生成研究画像</button>
    <div id="profile_box" class="hidden mb-4">
      <label class="block text-sm font-medium mb-1 text-orange-600">【可手动修改】研究画像：</label>
      <textarea id="t_profile" class="w-full border p-2 rounded h-32 text-sm"></textarea>
      <button onclick="startTxt()" class="bg-blue-600 text-white px-4 py-2 rounded mt-2">2. 确认无误，开始生成 TXT</button>
    </div>
    <div id="txt_progress" class="hidden"><div id="txt_msg" class="text-sm mb-2"></div><div class="w-full bg-gray-200 rounded h-3"><div id="txt_bar" class="bg-blue-600 h-3 rounded" style="width:0%"></div></div></div>
    <div id="txt_dl" class="hidden mt-4 p-4 bg-green-50 rounded text-center">
      <div class="text-green-700 font-bold mb-2">✅ TXT生成完成！请下载并检查。</div>
      <a id="a_txt" href="#" class="bg-blue-600 text-white px-6 py-2 rounded">📄 下载 00_完整论文.txt</a>
    </div>
    <div id="txt_queue" class="hidden mt-4 p-3 bg-red-50 text-red-700 font-bold rounded">⚠️ 系统繁忙，正在排队...</div>
  </div>

  <!-- 第二步：转DOCX -->
  <div class="bg-white p-6 rounded-lg shadow border-t-4 border-green-500">
    <h2 class="text-xl font-bold mb-4">第二步：上传 TXT 转换排版 DOCX</h2>
    <input type="file" id="f_upload" accept=".txt" class="mb-4 block w-full text-sm text-gray-500 file:mr-4 file:py-2 file:px-4 file:rounded file:border-0 file:text-sm file:font-semibold file:bg-blue-50 file:text-blue-700 hover:file:bg-blue-100">
    <button onclick="startDocx()" class="bg-green-600 text-white px-4 py-2 rounded disabled:opacity-50" id="btn_docx" disabled>开始转换 DOCX</button>
    <div id="docx_progress" class="hidden mt-4"><div id="docx_msg" class="text-sm mb-2"></div><div class="w-full bg-gray-200 rounded h-3"><div id="docx_bar" class="bg-green-600 h-3 rounded" style="width:0%"></div></div></div>
    <div id="docx_dl" class="hidden mt-4 p-4 bg-green-50 rounded text-center">
      <div class="text-green-700 font-bold mb-2">✅ 排版完成！</div>
      <a id="a_docx" href="#" class="bg-green-600 text-white px-6 py-2 rounded">📘 下载 论文_含图表.docx</a>
    </div>
    <div id="docx_queue" class="hidden mt-4 p-3 bg-red-50 text-red-700 font-bold rounded">⚠️ 系统繁忙，正在排队...</div>
  </div>
</div>

<script>
let tid1='', tid2='', iv1='', iv2='';
document.getElementById('f_upload').addEventListener('change', e => { document.getElementById('btn_docx').disabled = !e.target.files.length; });

async function req(url, body) { const r = await fetch(url, {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify(body)}); return r.json(); }
async function poll(id, type) {
    const d = await (await fetch(`/api/status/${id}`)).json();
    document.getElementById(`${type}_msg`).innerText = d.msg;
    document.getElementById(`${type}_bar`).style.width = d.progress+'%';
    if(d.status==='completed'){
        clearInterval(type==='txt' ? iv1 : iv2);
        document.getElementById(`${type}_progress`).classList.add('hidden');
        document.getElementById(`${type}_dl`).classList.remove('hidden');
        document.getElementById(`a_${type}`).href = `/download/${id}/${type==='txt'?'00_完整论文.txt':'论文_含图表.docx'}`;
    } else if(d.status==='error'){ clearInterval(type==='txt' ? iv1 : iv2); document.getElementById(`${type}_msg`).innerText="❌"+d.msg; }
}

async function genProfile(){
    const p = await req('/api/gen_profile', {title: t_title.value, major: t_major.value, company: t_comp.value, industry: t_ind.value});
    document.getElementById('t_profile').value = JSON.stringify(p, null, 2);
    document.getElementById('profile_box').classList.remove('hidden');
}

async function startTxt(){
    const profile = JSON.parse(document.getElementById('t_profile').value);
    const r = await req('/api/start_gen_txt', profile);
    if(r.status==='queue'){ document.getElementById('txt_queue').classList.remove('hidden'); iv1=setInterval(()=>startTxt(),3000); return; }
    tid1=r.task_id;
    document.getElementById('txt_queue').classList.add('hidden');
    document.getElementById('profile_box').classList.add('hidden');
    document.getElementById('txt_progress').classList.remove('hidden');
    iv1=setInterval(()=>poll(tid1,'txt'),1000);
}

async function startDocx(){
    const fd = new FormData(); fd.append('file', document.getElementById('f_upload').files[0]);
    const r = await fetch('/api/start_convert', {method:'POST', body:fd}).then(r=>r.json());
    if(r.status==='queue'){ document.getElementById('docx_queue').classList.remove('hidden'); iv2=setInterval(()=>startDocx(),3000); return; }
    tid2=r.task_id;
    document.getElementById('docx_queue').classList.add('hidden');
    document.getElementById('docx_progress').classList.remove('hidden');
    iv2=setInterval(()=>poll(tid2,'docx'),1000);
}
</script></body></html>"""

# ---------- 设计类页面 ----------
@app.get("/design", response_class=HTMLResponse)
async def design_page():
    return """
<!DOCTYPE html><html lang="zh-CN"><head><meta charset="UTF-8"><title>设计类论文生成</title><script src="https://cdn.tailwindcss.com"></script></head>
<body class="bg-gray-100 p-8">
<div class="max-w-4xl mx-auto space-y-8">
  <div class="flex items-center justify-between">
    <a href="/" class="text-blue-600 hover:underline">← 返回首页</a>
    <h1 class="text-3xl font-bold text-center text-purple-700">🎨 设计类毕业论文生成</h1>
    <div></div>
  </div>
  
  <!-- 第一步：生成TXT -->
  <div class="bg-white p-6 rounded-lg shadow border-t-4 border-purple-500">
    <h2 class="text-xl font-bold mb-4">第一步：生成论文 TXT 原文</h2>
    <div class="grid grid-cols-2 gap-4 mb-4">
      <div><label class="block text-sm font-medium mb-1">论文题目</label><input id="d_title" class="w-full border p-2 rounded" value="新中式风格女装系列设计研究"></div>
      <div><label class="block text-sm font-medium mb-1">设计专业</label><input id="d_type" class="w-full border p-2 rounded" value="服装设计"></div>
      <div><label class="block text-sm font-medium mb-1">设计对象</label><input id="d_obj" class="w-full border p-2 rounded" value="新中式女装"></div>
      <div><label class="block text-sm font-medium mb-1">使用场景</label><input id="d_ctx" class="w-full border p-2 rounded" value="日常穿着与礼仪场合"></div>
    </div>
    <button onclick="genDesignProfile()" class="bg-gray-500 text-white px-4 py-2 rounded mb-4">1. 自动生成设计画像</button>
    <div id="profile_box" class="hidden mb-4">
      <label class="block text-sm font-medium mb-1 text-orange-600">【可手动修改】设计画像：</label>
      <textarea id="d_profile" class="w-full border p-2 rounded h-32 text-sm"></textarea>
      <button onclick="startDesignTxt()" class="bg-purple-600 text-white px-4 py-2 rounded mt-2">2. 确认无误，开始生成 TXT</button>
    </div>
    <div id="txt_progress" class="hidden"><div id="txt_msg" class="text-sm mb-2"></div><div class="w-full bg-gray-200 rounded h-3"><div id="txt_bar" class="bg-purple-600 h-3 rounded" style="width:0%"></div></div></div>
    <div id="txt_dl" class="hidden mt-4 p-4 bg-green-50 rounded text-center">
      <div class="text-green-700 font-bold mb-2">✅ TXT生成完成！请下载并检查。</div>
      <a id="a_txt" href="#" class="bg-purple-600 text-white px-6 py-2 rounded">📄 下载 00_完整论文.txt</a>
    </div>
    <div id="txt_queue" class="hidden mt-4 p-3 bg-red-50 text-red-700 font-bold rounded">⚠️ 系统繁忙，正在排队...</div>
  </div>

  <!-- 第二步：转DOCX -->
  <div class="bg-white p-6 rounded-lg shadow border-t-4 border-green-500">
    <h2 class="text-xl font-bold mb-4">第二步：上传 TXT 转换排版 DOCX</h2>
    <input type="file" id="f_upload" accept=".txt" class="mb-4 block w-full text-sm text-gray-500 file:mr-4 file:py-2 file:px-4 file:rounded file:border-0 file:text-sm file:font-semibold file:bg-purple-50 file:text-purple-700 hover:file:bg-purple-100">
    <button onclick="startDocx()" class="bg-green-600 text-white px-4 py-2 rounded disabled:opacity-50" id="btn_docx" disabled>开始转换 DOCX</button>
    <div id="docx_progress" class="hidden mt-4"><div id="docx_msg" class="text-sm mb-2"></div><div class="w-full bg-gray-200 rounded h-3"><div id="docx_bar" class="bg-green-600 h-3 rounded" style="width:0%"></div></div></div>
    <div id="docx_dl" class="hidden mt-4 p-4 bg-green-50 rounded text-center">
      <div class="text-green-700 font-bold mb-2">✅ 排版完成！</div>
      <a id="a_docx" href="#" class="bg-green-600 text-white px-6 py-2 rounded">📘 下载 论文_含图表.docx</a>
    </div>
    <div id="docx_queue" class="hidden mt-4 p-3 bg-red-50 text-red-700 font-bold rounded">⚠️ 系统繁忙，正在排队...</div>
  </div>
</div>

<script>
let tid1='', tid2='', iv1='', iv2='';
document.getElementById('f_upload').addEventListener('change', e => { document.getElementById('btn_docx').disabled = !e.target.files.length; });

async function req(url, body) { const r = await fetch(url, {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify(body)}); return r.json(); }
async function poll(id, type) {
    const d = await (await fetch(`/api/status/${id}`)).json();
    document.getElementById(`${type}_msg`).innerText = d.msg;
    document.getElementById(`${type}_bar`).style.width = d.progress+'%';
    if(d.status==='completed'){
        clearInterval(type==='txt' ? iv1 : iv2);
        document.getElementById(`${type}_progress`).classList.add('hidden');
        document.getElementById(`${type}_dl`).classList.remove('hidden');
        document.getElementById(`a_${type}`).href = `/download/${id}/${type==='txt'?'00_完整论文.txt':'论文_含图表.docx'}`;
    } else if(d.status==='error'){ clearInterval(type==='txt' ? iv1 : iv2); document.getElementById(`${type}_msg`).innerText="❌"+d.msg; }
}

async function genDesignProfile(){
    const p = await req('/api/gen_design_profile', {title: d_title.value, design_type: d_type.value, design_object: d_obj.value, context: d_ctx.value});
    document.getElementById('d_profile').value = JSON.stringify(p, null, 2);
    document.getElementById('profile_box').classList.remove('hidden');
}

async function startDesignTxt(){
    const profile = JSON.parse(document.getElementById('d_profile').value);
    const r = await req('/api/start_gen_design_txt', profile);
    if(r.status==='queue'){ document.getElementById('txt_queue').classList.remove('hidden'); iv1=setInterval(()=>startDesignTxt(),3000); return; }
    tid1=r.task_id;
    document.getElementById('txt_queue').classList.add('hidden');
    document.getElementById('profile_box').classList.add('hidden');
    document.getElementById('txt_progress').classList.remove('hidden');
    iv1=setInterval(()=>poll(tid1,'txt'),1000);
}

async function startDocx(){
    const fd = new FormData(); fd.append('file', document.getElementById('f_upload').files[0]);
    const r = await fetch('/api/start_convert', {method:'POST', body:fd}).then(r=>r.json());
    if(r.status==='queue'){ document.getElementById('docx_queue').classList.remove('hidden'); iv2=setInterval(()=>startDocx(),3000); return; }
    tid2=r.task_id;
    document.getElementById('docx_queue').classList.add('hidden');
    document.getElementById('docx_progress').classList.remove('hidden');
    iv2=setInterval(()=>poll(tid2,'docx'),1000);
}
</script></body></html>"""

# ================== [ API 路由 ] ==================

# ---------- 经管类 API ----------
@app.post("/api/gen_profile")
async def api_gen_profile(request: Request):
    data = await request.json()
    title = data.get("title")
    major = data.get("major")
    company = data.get("company")
    industry = data.get("industry")
    if not title: title = "请在网页输入论文题目"
    if not major: major = "请在网页输入专业"
    if not company: company = "请在网页输入公司"
    if not industry: industry = "请在网页输入行业"
    print(f"✅ [经管类-收到参数] 准备调LLM: {title} / {major} / {company}")
    profile = generate_profile_web(major, title, company)
    if not profile.get("core_problems"):
        profile["core_problems"] = ["等待补充问题1", "等待补充问题2"]
    if not profile.get("data_hints"):
        profile["data_hints"] = ["等待补充数据1", "等待补充数据2"]
    return profile

@app.post("/api/start_gen_txt")
async def api_start_txt(request: Request):
    global is_system_busy, system_lock
    data = await request.json()
    with system_lock:
        if is_system_busy: return {"status": "queue"}
        is_system_busy = True
        tid = str(uuid.uuid4())
        tasks_db[tid] = {"status": "running", "progress": 0, "msg": "初始化...", "files": []}
        threading.Thread(target=run_txt_pipeline, args=(tid, data)).start()
        return {"status": "ok", "task_id": tid}

# ---------- 设计类 API ----------
@app.post("/api/gen_design_profile")
async def api_gen_design_profile(request: Request):
    data = await request.json()
    title = data.get("title")
    design_type = data.get("design_type")
    design_object = data.get("design_object")
    context = data.get("context", "")
    if not title: title = "请在网页输入论文题目"
    if not design_type: design_type = "请在网页输入设计专业"
    if not design_object: design_object = "请在网页输入设计对象"
    print(f"✅ [设计类-收到参数] 准备调LLM: {title} / {design_type} / {design_object}")
    profile = generate_design_profile(design_type, title, design_object)
    if context:
        profile["context"] = context
    if not profile.get("core_problems"):
        profile["core_problems"] = ["等待补充问题1", "等待补充问题2", "等待补充问题3"]
    if not profile.get("design_strategies"):
        profile["design_strategies"] = ["等待补充策略1", "等待补充策略2", "等待补充策略3"]
    return profile

@app.post("/api/start_gen_design_txt")
async def api_start_design_txt(request: Request):
    global is_system_busy, system_lock
    data = await request.json()
    with system_lock:
        if is_system_busy: return {"status": "queue"}
        is_system_busy = True
        tid = str(uuid.uuid4())
        tasks_db[tid] = {"status": "running", "progress": 0, "msg": "初始化...", "files": []}
        threading.Thread(target=run_design_txt_pipeline, args=(tid, data)).start()
        return {"status": "ok", "task_id": tid}

# ---------- 公共 API ----------
@app.post("/api/start_convert")
async def api_start_convert(file: UploadFile = File(...)):
    global is_system_busy, system_lock
    with system_lock:
        if is_system_busy: return {"status": "queue"}
        is_system_busy = True
        tid = str(uuid.uuid4())
        tasks_db[tid] = {"status": "running", "progress": 0, "msg": "接收文件...", "files": []}
        temp_path = f"temp_{tid}.txt"
        with open(temp_path, "wb") as f: f.write(await file.read())
        threading.Thread(target=run_docx_pipeline, args=(tid, temp_path)).start()
        return {"status": "ok", "task_id": tid}

@app.get("/api/status/{task_id}")
async def api_status(task_id: str):
    return tasks_db.get(task_id, {"status": "not_found", "msg": "未知", "progress": 0})

@app.get("/download/{task_id}/{filename}")
async def api_download(task_id: str, filename: str):
    path = os.path.join("output", task_id, filename)
    if os.path.exists(path): return FileResponse(path, filename=filename)
    return {"error": "not found"}

if __name__ == "__main__":
    import uvicorn
    print("\n🚀 启动成功: http://localhost:8000\n")
    print("经管类论文: http://localhost:8000/manage")
    print("设计类论文: http://localhost:8000/design\n")
    uvicorn.run(app, host="0.0.0.0", port=8000)
